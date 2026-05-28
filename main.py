import os
import uuid
import shutil
import json
import base64
from pathlib import Path
from typing import Optional

import anthropic
import fitz
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import uvicorn

app = FastAPI(title="File Converter API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["X-Original-Size", "X-Compressed-Size", "X-Reduction"],
)

UPLOAD_DIR = Path("/tmp/uploads")
OUTPUT_DIR = Path("/tmp/outputs")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

IMAGE_EXTS = {"png", "jpg", "jpeg", "gif", "webp", "bmp", "tiff"}

def save_upload(file: UploadFile) -> Path:
    ext  = Path(file.filename).suffix.lower()
    dest = UPLOAD_DIR / f"{uuid.uuid4()}{ext}"
    with dest.open("wb") as f:
        shutil.copyfileobj(file.file, f)
    return dest

def cleanup(*paths: Path):
    for p in paths:
        try:
            if p and p.exists(): p.unlink()
        except: pass

def is_image_ext(ext: str) -> bool:
    return ext.lstrip(".") in IMAGE_EXTS


# ── Health check ───────────────────────────────────────────────────────────────

@app.get("/")
async def health():
    return {"status": "ok", "service": "File Converter API"}


# ── AI: phân tích file ─────────────────────────────────────────────────────────

@app.post("/api/ai/process")
async def ai_process(
    file: UploadFile = File(...),
    tool: str = Form(...),
    options: Optional[str] = Form(None),
):
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(500, "ANTHROPIC_API_KEY chưa được cấu hình")

    opts = json.loads(options) if options else {}
    path = save_upload(file)
    ext  = path.suffix.lower()

    try:
        client = anthropic.Anthropic(api_key=api_key)
        with open(path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()

        media_map = {
            ".pdf": "application/pdf", ".png": "image/png",
            ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".gif": "image/gif", ".webp": "image/webp",
        }
        media_type = media_map.get(ext, "application/pdf")
        is_image   = is_image_ext(ext)

        if tool == "extract":
            prompt = "Extract all text from this document. Preserve headings, paragraphs, lists, and tables."
        elif tool == "translate":
            lang = opts.get("lang", "vi")
            lang_names = {
                "vi":"Vietnamese (tiếng Việt)", "en":"English",
                "ja":"Japanese (日本語)", "zh":"Chinese (中文)",
                "fr":"French (Français)", "ko":"Korean (한국어)",
                "de":"German (Deutsch)", "es":"Spanish (Español)",
            }
            prompt = f"Translate all text to {lang_names.get(lang, lang)}. Return only translated text."
        elif tool == "summarize":
            fmt = opts.get("format", "bullet")
            prompt = ("Summarize as clear bullet points." if fmt == "bullet"
                      else "Write a concise paragraph summary (3–5 sentences).")
        elif tool == "qa":
            question = opts.get("question", "What is this document about?")
            prompt = f"Answer this question based on the document:\n\n{question}"
        elif tool == "ocr":
            prompt = "Perform OCR. Extract ALL visible text exactly as it appears."
        else:
            raise HTTPException(400, f"Tool không hợp lệ: {tool}")

        content = (
            [{"type":"image",    "source":{"type":"base64","media_type":media_type,        "data":b64}}, {"type":"text","text":prompt}]
            if is_image else
            [{"type":"document", "source":{"type":"base64","media_type":"application/pdf", "data":b64}}, {"type":"text","text":prompt}]
        )

        response = client.messages.create(
            model="claude-sonnet-4-5", max_tokens=4096,
            messages=[{"role":"user","content":content}],
        )
        result_text = "".join(b.text for b in response.content if hasattr(b, "text"))
        return {"success": True, "result": result_text}
    finally:
        cleanup(path)


# ── PDF → Word ─────────────────────────────────────────────────────────────────

@app.post("/api/convert/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Chỉ chấp nhận file PDF")

    path     = save_upload(file)
    out_path = OUTPUT_DIR / f"{uuid.uuid4()}.docx"

    try:
        from pdf2docx import Converter
        cv = Converter(str(path))
        cv.convert(str(out_path), start=0, end=None)
        cv.close()

        if not out_path.exists():
            raise HTTPException(500, "Không tạo được file DOCX")

        return FileResponse(
            str(out_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=Path(file.filename).stem + ".docx",
        )
    finally:
        cleanup(path)


# ── Word → PDF (dùng python-docx + reportlab, không cần LibreOffice) ──────────

@app.post("/api/convert/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    ext = Path(file.filename).suffix.lower()
    if ext not in {".doc", ".docx"}:
        raise HTTPException(400, "Chỉ chấp nhận .doc hoặc .docx")

    path     = save_upload(file)
    out_path = OUTPUT_DIR / f"{uuid.uuid4()}.pdf"

    try:
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import re

        # Đọc nội dung DOCX
        doc = Document(str(path))

        # Tạo PDF
        pdf_doc = SimpleDocTemplate(
            str(out_path),
            pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm,
        )

        styles    = getSampleStyleSheet()
        story     = []

        # Style tùy chỉnh
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=6,
        )
        h1_style = ParagraphStyle(
            'CustomH1',
            parent=styles['Heading1'],
            fontSize=16,
            leading=20,
            spaceAfter=10,
            spaceBefore=14,
        )
        h2_style = ParagraphStyle(
            'CustomH2',
            parent=styles['Heading2'],
            fontSize=13,
            leading=18,
            spaceAfter=8,
            spaceBefore=10,
        )

        def clean_text(text):
            # Escape các ký tự đặc biệt trong ReportLab
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return text

        # Duyệt qua các paragraph trong DOCX
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 6))
                continue

            style_name = para.style.name.lower()

            if 'heading 1' in style_name:
                story.append(Paragraph(clean_text(text), h1_style))
            elif 'heading 2' in style_name:
                story.append(Paragraph(clean_text(text), h2_style))
            else:
                story.append(Paragraph(clean_text(text), normal_style))

        # Duyệt qua các bảng trong DOCX
        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = [Paragraph(clean_text(cell.text), normal_style) for cell in row.cells]
                data.append(row_data)

            if data:
                t = Table(data, repeatRows=1)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#E8E3FF')),
                    ('TEXTCOLOR',  (0,0), (-1,0), colors.HexColor('#2E2B45')),
                    ('GRID',       (0,0), (-1,-1), 0.5, colors.HexColor('#E4E1F0')),
                    ('FONTSIZE',   (0,0), (-1,-1), 10),
                    ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F8F7FC')]),
                    ('TOPPADDING',  (0,0), (-1,-1), 6),
                    ('BOTTOMPADDING',(0,0),(-1,-1), 6),
                    ('LEFTPADDING', (0,0), (-1,-1), 8),
                    ('RIGHTPADDING',(0,0), (-1,-1), 8),
                ]))
                story.append(Spacer(1, 8))
                story.append(t)
                story.append(Spacer(1, 8))

        if not story:
            story.append(Paragraph("(Tài liệu trống)", normal_style))

        pdf_doc.build(story)

        if not out_path.exists():
            raise HTTPException(500, "Không tạo được file PDF")

        return FileResponse(
            str(out_path),
            media_type="application/pdf",
            filename=Path(file.filename).stem + ".pdf",
        )

    except ImportError as e:
        raise HTTPException(500, f"Thiếu thư viện: {str(e)}")
    finally:
        cleanup(path)


# ── PDF → Image ────────────────────────────────────────────────────────────────

@app.post("/api/convert/pdf-to-image")
async def pdf_to_image(
    file: UploadFile = File(...),
    page: int = Form(0),
    dpi:  int = Form(150),
    fmt:  str = Form("png"),
):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Chỉ chấp nhận file PDF")

    path     = save_upload(file)
    out_path = OUTPUT_DIR / f"{uuid.uuid4()}.{fmt}"

    try:
        doc = fitz.open(str(path))
        if page >= len(doc):
            raise HTTPException(400, f"PDF chỉ có {len(doc)} trang")
        pix = doc[page].get_pixmap(matrix=fitz.Matrix(dpi/72, dpi/72))
        pix.save(str(out_path))
        doc.close()
        return FileResponse(
            str(out_path),
            media_type="image/png" if fmt=="png" else "image/jpeg",
            filename=f"{Path(file.filename).stem}_page{page+1}.{fmt}",
        )
    finally:
        cleanup(path)


# ── Merge PDF ──────────────────────────────────────────────────────────────────

@app.post("/api/convert/merge-pdf")
async def merge_pdf(files: list[UploadFile] = File(...)):
    if len(files) < 2:
        raise HTTPException(400, "Cần ít nhất 2 file PDF")

    paths = []; out_path = OUTPUT_DIR / f"{uuid.uuid4()}.pdf"
    try:
        for f in files:
            if not f.filename.lower().endswith(".pdf"):
                raise HTTPException(400, f"'{f.filename}' không phải PDF")
            paths.append(save_upload(f))
        merged = fitz.open()
        for p in paths:
            doc = fitz.open(str(p)); merged.insert_pdf(doc); doc.close()
        merged.save(str(out_path)); merged.close()
        return FileResponse(str(out_path), media_type="application/pdf", filename="merged.pdf")
    finally:
        for p in paths: cleanup(p)


# ── Split PDF ──────────────────────────────────────────────────────────────────

@app.post("/api/convert/split-pdf")
async def split_pdf(file: UploadFile = File(...), pages: str = Form(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Chỉ chấp nhận file PDF")

    path = save_upload(file); out_path = OUTPUT_DIR / f"{uuid.uuid4()}.pdf"
    try:
        doc = fitz.open(str(path)); total = len(doc)
        page_nums = set()
        for part in pages.split(","):
            part = part.strip()
            if not part: continue
            if "-" in part:
                s, e = part.split("-", 1)
                for n in range(int(s), int(e)+1): page_nums.add(n-1)
            else: page_nums.add(int(part)-1)
        invalid = [n+1 for n in page_nums if n < 0 or n >= total]
        if invalid: raise HTTPException(400, f"Trang không tồn tại: {invalid}")
        new_doc = fitz.open()
        for n in sorted(page_nums): new_doc.insert_pdf(doc, from_page=n, to_page=n)
        new_doc.save(str(out_path)); new_doc.close(); doc.close()
        return FileResponse(str(out_path), media_type="application/pdf",
            filename=f"{Path(file.filename).stem}_split.pdf")
    finally:
        cleanup(path)


# ── Compress PDF ───────────────────────────────────────────────────────────────

@app.post("/api/convert/compress-pdf")
async def compress_pdf(file: UploadFile = File(...), level: str = Form("medium")):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Chỉ chấp nhận file PDF")

    path = save_upload(file); out_path = OUTPUT_DIR / f"{uuid.uuid4()}.pdf"
    try:
        # Dùng PyMuPDF để nén — không cần Ghostscript
        doc = fitz.open(str(path))
        # Garbage collect và deflate để giảm kích thước
        doc.save(
            str(out_path),
            garbage=4,       # xóa object thừa
            deflate=True,    # nén stream
            clean=True,      # clean content streams
        )
        doc.close()

        orig = path.stat().st_size
        comp = out_path.stat().st_size
        reduction = round((1 - comp / orig) * 100, 1)

        resp = FileResponse(
            str(out_path), media_type="application/pdf",
            filename=f"{Path(file.filename).stem}_compressed.pdf",
        )
        resp.headers["X-Original-Size"]   = str(orig)
        resp.headers["X-Compressed-Size"] = str(comp)
        resp.headers["X-Reduction"]       = str(reduction)
        return resp
    finally:
        cleanup(path)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("main:app", host="0.0.0.0", port=port)
